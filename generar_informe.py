# -*- coding: utf-8 -*-
"""
Genera el Pre Informe de Tesis basado en el template de la escuela,
con los resultados reales del proyecto (Estrategia A, no B).
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

# ─── Abrir template ──────────────────────────────────────────────────────────
doc = Document('Pre informe de tesis - 2026-I.docx')
body = doc.element.body

# ─── Eliminar todo desde el heading RESULTADOS en adelante ───────────────────
children = list(body)
remove_from = None
for i, child in enumerate(children):
    if child.tag == qn('w:p'):
        pPr = child.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                val = pStyle.get(qn('w:val'), '')
                # Template uses Spanish-localized style ID 'Ttulo1' for Heading 1
                if val in ('Ttulo1', 'Heading1', '1'):
                    txt = ''.join(t.text or '' for t in child.findall('.//' + qn('w:t')))
                    if 'RESULTADOS' in txt.upper():
                        remove_from = i
                        break

if remove_from is not None:
    to_remove = [c for c in children[remove_from:] if c.tag != qn('w:sectPr')]
    for el in to_remove:
        body.remove(el)

# ─── Helpers ──────────────────────────────────────────────────────────────────
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, style='Normal'):
    return doc.add_paragraph(text, style=style)

def bullet(text):
    p = doc.add_paragraph(style='List Paragraph')
    p.add_run(text)
    return p

def bold_intro(label, rest):
    """Párrafo con inicio en negrita seguido de texto normal."""
    p = doc.add_paragraph(style='Normal')
    run_b = p.add_run(label)
    run_b.bold = True
    p.add_run(rest)
    return p

def add_caption(text):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.italic = True
    return p

def tabla(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Cabecera
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    # Datos
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i + 1].cells[j].text = str(val)
    return t

def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

# ═══════════════════════════════════════════════════════════════════════════════
# I. RESULTADOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('RESULTADOS')

para(
    'Los resultados presentados a continuación corresponden a las cuatro primeras fases '
    'de la metodología CRISP-DM ejecutadas durante el Seminario de Tesis I, y al '
    'cumplimiento de los cuatro objetivos específicos trabajados hasta la fecha. '
    'La investigación combina CRISP-DM para la gestión analítica del proceso de modelado [1] '
    'con SCRUM —previsto para la fase de despliegue (Backend y Frontend)— que se encuentra '
    'en planificación.'
)

# ───────────────────────────────────────────────────────────────────────────────
# 1.1  En base a la metodología utilizada
# ───────────────────────────────────────────────────────────────────────────────
h2('En base a la metodología utilizada')

para(
    'El proceso de desarrollo siguió las fases de la metodología CRISP-DM: '
    'comprensión del negocio, comprensión de los datos, preparación de los datos y modelado. '
    'Cada fase se describe a continuación como una iteración del proceso de investigación.'
)

# ── 1.1.1 ─────────────────────────────────────────────────────────────────────
h3('Iteración #1: Comprensión del negocio')

para(
    'Se analizó el problema de la somnolencia y fatiga mental en el personal sanitario del '
    'consultorio oftalmológico NOR VISIÓN (Chiclayo, Perú). Se identificó el pluriempleo '
    'médico normalizado —habilitado por la Ley N.° 32145 (2024)— como causa raíz del riesgo '
    'de negligencias. Se definió el objetivo general: desarrollar un sistema embebido '
    'multimodal para la auto-detección de somnolencia y fatiga mental en cirujanos para '
    'prevenir negligencias médicas.'
)
para(
    'Se estableció la arquitectura de la solución en tres módulos: '
    '(i) Módulo 1 — visión conductual basado en aprendizaje profundo; '
    '(ii) Módulo 2 — fisiológico basado en motor de reglas con baseline personal; y '
    '(iii) Módulo 3 — fusión tardía que integra las probabilidades de ambos módulos '
    'mediante ponderación 40 % visión / 60 % fisiológico, con regla OR para señales de '
    'alta confianza (> 0.85 en cualquiera de los dos módulos).'
)
para(
    'Se definieron 18 requerimientos funcionales (RF-01 a RF-18) y 12 requerimientos '
    'no funcionales (RNF-01 a RNF-12). Los más relevantes para la validación experimental son:'
)
bullet('RF-06: El sistema retornará P_somnolencia ∈ [0, 1] por frame de evaluación.')
bullet('RNF-01: La evaluación completa (Módulo 1 + 2 + 3) se procesará en ≤ 10 segundos.')
bullet('RNF-02: El modelo de visión deberá alcanzar exactitud mínima de 80 % sobre el '
       'conjunto de prueba del dataset verificado.')
bullet('RNF-04: El dictamen mostrará explícitamente P_somnolencia, P_fatiga_fisiológica '
       'y P_total para garantizar la interpretabilidad.')

# ── 1.1.2 ─────────────────────────────────────────────────────────────────────
h3('Iteración #2: Comprensión y análisis de los datos')

para(
    'Se seleccionó el Driver Drowsiness Dataset (DDD) [2] como dataset principal para '
    'el Módulo 1. El dataset contiene 41 793 imágenes faciales RGB (resolución 227×227 px) '
    'de 28 sujetos únicos (identificados A–ZC), con la siguiente distribución de clases: '
    '19 445 imágenes alert y 22 348 imágenes drowsy. Cada sujeto contribuye con un promedio '
    'de aproximadamente 1 492 imágenes, lo que introduce un riesgo metodológico de data '
    'leakage si la partición no controla la identidad del sujeto.'
)
para(
    'Se realizó un análisis estadístico de las variables conductuales EAR (Eye Aspect Ratio) '
    'y MAR (Mouth Aspect Ratio) extraídas mediante MediaPipe FaceMesh sobre n = 41 776 frames. '
    'Los resultados se reportan en el OE-01 (sección 1.2.1). Se identificaron y sustentaron '
    'en la literatura cuatro variables fisiológicas (RMS-EMG, frecuencia mediana EMG, HR y '
    'HRV), completando el conjunto de ocho variables requerido por el indicador del OE-01.'
)

# ── 1.1.3 ─────────────────────────────────────────────────────────────────────
h3('Iteración #3: Preparación de los datos y estrategias de partición')

para(
    'Se implementaron y documentaron dos estrategias de partición del dataset. Esta decisión '
    'metodológica dual constituye una contribución central del trabajo, ya que permite '
    'documentar explícitamente la brecha entre evaluación honesta y evaluación inflada.'
)
bold_intro(
    'Estrategia A — Subject-independent split (partición por sujeto): ',
    'Los 28 sujetos se repartieron de forma disjunta: 19 sujetos para entrenamiento '
    '(26 775 imágenes), 4 sujetos para validación (6 404 imágenes) y 5 sujetos para '
    'prueba (8 614 imágenes). Esta estrategia garantiza que ningún frame del sujeto '
    'evaluado aparece en el entrenamiento, replicando el escenario real de NOR VISIÓN '
    'donde cada médico evaluado es nuevo para el modelo [3]. Es el esquema de evaluación '
    'más cercano al despliegue real y el adoptado como referencia científica en este trabajo.'
)
bold_intro(
    'Estrategia B — Split aleatorio por imagen (referencia comparativa con la literatura): ',
    'Las 41 793 imágenes se mezclaron y distribuyeron aleatoriamente con estratificación '
    'por clase (entrenamiento: 29 255; validación: 6 269; prueba: 6 269). Este esquema '
    'replica el procedimiento estándar de la literatura sobre DDD [2], permitiendo '
    'comparación directa con trabajos previos publicados. Sin embargo, introduce data '
    'leakage implícito: imágenes del mismo sujeto aparecen simultáneamente en entrenamiento '
    'y prueba, permitiendo al modelo memorizar rasgos biométricos individuales.'
)
para(
    'Para el módulo BiLSTM, los valores de EAR y MAR se extrajeron frame a frame mediante '
    'MediaPipe FaceMesh (versión 0.10.14, la última estable compatible). Las secuencias '
    'se construyeron con sliding window (SEQ_LEN = 20, STRIDE = 5) por sujeto (Estrategia A) '
    'o aleatoriamente (Estrategia B). La normalización se realizó exclusivamente con los '
    'estadísticos del conjunto de entrenamiento, sin data leakage en la escala.'
)

# ── 1.1.4 ─────────────────────────────────────────────────────────────────────
h3('Iteración #4: Modelado')

para(
    'Se entrenaron tres arquitecturas representativas del estado del arte en visión '
    'artificial para la detección de somnolencia [4][5][6][7], bajo ambas estrategias '
    'de partición:'
)
bold_intro(
    'MobileNetV2 (CNN): ',
    'Backbone pre-entrenado en ImageNet-1k con fine-tuning en dos fases: '
    'Fase 1 (8 épocas, lr = 1×10⁻³, solo cabeza clasificadora) y Fase 2 '
    '(18 épocas, lr = 3×10⁻⁵, fine-tuning completo). Optimizador AdamW, '
    'cosine LR scheduler, label smoothing 0.1, class weights y mixed precision training (AMP).'
)
bold_intro(
    'BiLSTM sobre secuencias de landmarks faciales: ',
    'Arquitectura BiLSTM de 2 capas (hidden = 128) seguida de pooling combinado '
    '(mean + last token) y MLP clasificador. Entrenado sobre secuencias de EAR+MAR '
    'extraídas por MediaPipe FaceMesh, con normalización sin data leakage.'
)
bold_intro(
    'ViT-Tiny (Vision Transformer): ',
    'Modelo pre-entrenado en ImageNet vía timm. Optimizador AdamW '
    '(lr = 3×10⁻⁴, weight decay = 5×10⁻⁴). Warmup lineal (3 épocas) + cosine decay, '
    'hasta 25 épocas con early stopping por F1 de validación [6].'
)
para(
    'Para CNN y ViT se aplicó Test-Time Augmentation (TTA): se promediaron las '
    'probabilidades de la imagen original y su flip horizontal, reduciendo la varianza '
    'de la predicción. El mejor checkpoint de cada modelo se seleccionó por F1-score '
    'en validación.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 1.2  En base a los objetivos del proyecto
# ═══════════════════════════════════════════════════════════════════════════════
h2('En base a los objetivos del proyecto')

para(
    'A continuación se demuestra con evidencias el cumplimiento de los cuatro objetivos '
    'específicos trabajados en la fase actual de la investigación.'
)

# ── 1.2.1 OE-01 ───────────────────────────────────────────────────────────────
h3('OE-01: Determinar las variables relacionadas a los factores conductuales y '
   'fisiológicos asociados a la somnolencia y fatiga mental en cirujanos')

para(
    'Se identificaron y caracterizaron ocho variables relevantes para la detección de '
    'somnolencia y fatiga mental, superando el mínimo de seis exigido en el indicador '
    'del objetivo. Las variables se organizan en dos grupos:'
)
bold_intro('Variables conductuales (4): ', 'EAR, MAR, PERCLOS y pose de cabeza (pitch, yaw, roll).')
bold_intro('Variables fisiológicas (4): ',
           'RMS-EMG del trapecio superior, frecuencia mediana EMG, '
           'HR (frecuencia cardíaca) y HRV (SDNN, RMSSD, pNN50).')

para(
    'Las variables conductuales EAR y MAR fueron validadas estadísticamente mediante '
    'prueba t de Student de dos muestras independientes sobre n = 41 776 frames del DDD. '
    'Los resultados se presentan en la Tabla 1.'
)

add_caption('Tabla 1. Validación estadística de variables conductuales sobre DDD (n = 41 776 frames).')
tabla(
    ['Variable', 'Media Alerta', 'Media Somnoliento', 'Estadístico t', 'p-valor', "Cohen's d", 'Significativa'],
    [
        ['EAR (Eye Aspect Ratio)', '0.2718', '0.2386', '51.86', '≈ 0.00', '0.501 (mediano)', 'Sí'],
        ['MAR (Mouth Aspect Ratio)', '0.0376', '0.0324', '20.40', '5.49 × 10⁻⁹²', '0.195 (trivial)', 'Sí'],
    ]
)
doc.add_paragraph()

para(
    'Ambas variables presentan diferencias estadísticamente significativas entre los grupos '
    'alerta y somnoliento (p < 0.05). El EAR muestra efecto mediano (Cohen\'s d = 0.501): '
    'el cierre parcial de ojos es el discriminador más robusto de somnolencia. El MAR '
    'muestra efecto trivial (d = 0.195), pero la significancia estadística a nivel '
    'poblacional (p = 5.49 × 10⁻⁹²) lo confirma como variable complementaria. '
    'Las variables PERCLOS y pose de cabeza se validan en la literatura por '
    'Pathak et al. [5] y Zaman et al. [6]. Las variables fisiológicas se sustentan '
    'en Wijsman et al. [8] (RMS-EMG del trapecio como predictor de estrés cognitivo) '
    'y Zhang et al. [9] (fatiga mental y física en cirujanos laparoscópicos).'
)

# ── 1.2.2 OE-02 ───────────────────────────────────────────────────────────────
h3('OE-02: Comparar algoritmos basados en visión artificial para la selección del '
   'algoritmo de detección temprana de somnolencia y fatiga mental en cirujanos')

para(
    'Se compararon tres arquitecturas de aprendizaje profundo bajo dos estrategias de '
    'evaluación sobre el DDD. La Tabla 2 presenta los resultados completos. Se incluyen '
    'ambas estrategias para documentar la brecha metodológica identificada en la '
    'literatura, contribución científica de este trabajo.'
)

add_caption(
    'Tabla 2. Comparación de CNN, LSTM y ViT bajo Estrategia A (subject-independent) '
    'y Estrategia B (split aleatorio) sobre DDD.'
)
tabla(
    ['Modelo', 'Estrategia', 'Accuracy\n(%)', 'Precisión\n(%)', 'Sensibilidad\n(%)',
     'Especificidad\n(%)', 'F1\n(%)', 'AUC-ROC', 'Inferencia\n(ms)'],
    [
        ['MobileNetV2 (CNN)', 'A — subject-indep.', '53.05', '50.81', '79.01', '28.94', '61.85', '0.5135', '1.09'],
        ['BiLSTM (EAR+MAR)',  'A — subject-indep.', '74.56', '68.43', '87.59', '62.44', '76.83', '0.7942', '0.11'],
        ['ViT-Tiny',          'A — subject-indep.', '49.41', '47.91', '57.72', '41.68', '52.36', '0.5328', '1.08'],
        ['MobileNetV2 (CNN)', 'B — aleatorio',      '100.00','100.00','100.00','100.00','100.00','1.0000','1.12'],
        ['BiLSTM (EAR+MAR)',  'B — aleatorio',      '96.73', '96.85', '96.99', '96.43', '96.92', '0.9947', '0.15'],
        ['ViT-Tiny',          'B — aleatorio',      '99.92', '99.94', '99.91', '99.93', '99.93', '1.0000', '1.10'],
    ]
)
doc.add_paragraph()

para('Análisis de resultados:')
bullet(
    'Estrategia A (subject-independent): Los tres modelos alcanzan entre 49.4 % y 74.6 % '
    'de accuracy. El BiLSTM (EAR+MAR) resulta ganador con F1 = 76.83 % y '
    'sensibilidad = 87.59 %, superando a MobileNetV2 y ViT-Tiny en el escenario de '
    'evaluación honesta. Estos valores reflejan la dificultad real de generalizar a '
    'sujetos nunca vistos durante el entrenamiento, que es el escenario de NOR VISIÓN.'
)
bullet(
    'Estrategia B (split aleatorio): Los tres modelos alcanzan entre 96.7 % y 100.0 % '
    'de accuracy, comparables con los reportados en la literatura sobre DDD. Sin embargo, '
    'estas métricas están infladas por el data leakage implícito: el modelo aprende rasgos '
    'biométricos individuales de los sujetos en lugar de patrones generalizables de '
    'somnolencia. Por este motivo, los modelos de Estrategia B no se adoptan como '
    'modelos de producción.'
)
bullet(
    'Contribución metodológica: La brecha entre estrategias alcanza hasta 46.95 puntos '
    'porcentuales en accuracy para MobileNetV2 (100.00 % vs. 53.05 %). Este trabajo '
    'documenta explícitamente esta discrepancia como advertencia práctica: los resultados '
    'con split aleatorio típicamente reportados en la literatura (>95 % en DDD) '
    'sobreestiman significativamente la capacidad real de generalización a usuarios '
    'nuevos [3].'
)

# ── 1.2.3 OE-03 ───────────────────────────────────────────────────────────────
h3('OE-03: Generar el modelo de aprendizaje automático basado en el algoritmo '
   'seleccionado para la detección temprana de somnolencia y fatiga mental en cirujanos')

para(
    'El modelo seleccionado para el Módulo 1 del sistema es el '
    'BiLSTM (EAR+MAR), evaluado bajo la Estrategia A (subject-independent). '
    'La selección se fundamenta en: (i) mayor F1-score entre las tres arquitecturas '
    'en el escenario de evaluación honesta; (ii) mayor sensibilidad (87.59 %), '
    'criterio prioritario en aplicaciones médicas donde un falso negativo '
    '—no detectar somnolencia en un médico fatigado— tiene consecuencias más graves '
    'que un falso positivo; y (iii) latencia de inferencia de 0.11 ms/imagen, '
    'la más baja de las tres arquitecturas, garantizando ampliamente el RNF-01. '
    'Las métricas finales del modelo seleccionado se presentan en la Tabla 3.'
)

add_caption('Tabla 3. Métricas del modelo seleccionado: BiLSTM (EAR+MAR), Estrategia A — subject-independent.')
tabla(
    ['Métrica', 'Valor con umbral 0.50 (default)'],
    [
        ['Accuracy',              '74.56 %'],
        ['F1-score',              '76.83 %'],
        ['AUC-ROC',               '0.7942'],
        ['Sensibilidad (recall)', '87.59 %'],
        ['Especificidad',         '62.44 %'],
        ['Precisión',             '68.43 %'],
        ['Latencia de inferencia','0.11 ms/imagen'],
    ]
)
doc.add_paragraph()

para(
    'El modelo fue exportado en formato PyTorch (.pt) con pesos, metadatos de entrenamiento '
    'y estadísticos de normalización, listo para integración en el backend del sistema. '
    'La arquitectura es: BiLSTM (2 capas, hidden = 128) → pooling (mean + last token) '
    '→ MLP clasificador.'
)
para(
    'Respecto al RNF-02 (exactitud ≥ 80 %): con evaluación honesta (Estrategia A), '
    'el modelo alcanza 74.56 %, por debajo del umbral. Esta limitación se documenta '
    'honestamente; representa la dificultad intrínseca de la generalización a sujetos '
    'nuevos y constituye la contribución académica central de este trabajo. '
    'La sensibilidad de 87.59 % y la latencia de 0.11 ms/imagen cumplen los criterios '
    'de OE-04 y RNF-01. El ajuste fino del sistema en producción con datos de NOR VISIÓN '
    'queda diferido al OE-06 (integración con hardware embebido).'
)

# ── 1.2.4 OE-04 ───────────────────────────────────────────────────────────────
h3('OE-04: Evaluar la precisión, sensibilidad y especificidad del modelo de '
   'aprendizaje automático mediante métricas de desempeño estadístico')

para(
    'Se implementó una función de threshold tuning para identificar el umbral de '
    'decisión que permita cumplir simultáneamente los criterios del OE-04: '
    'precisión ≥ 90 %, sensibilidad ≥ 85 % y especificidad ≥ 85 %. La metodología '
    'evalúa 91 umbrales en el rango [0.05–0.95] sobre el conjunto de prueba '
    '(Estrategia A), seleccionando el umbral que maximiza F1-score entre los que '
    'satisfacen los tres criterios simultáneamente.'
)
para('Con el umbral por defecto (0.50), el BiLSTM (Estrategia A) obtiene:')
bullet('Sensibilidad: 87.59 % → supera el criterio de ≥ 85 %.')
bullet('Precisión: 68.43 % → por debajo del criterio de ≥ 90 %.')
bullet('Especificidad: 62.44 % → por debajo del criterio de ≥ 85 %.')
para(
    'El análisis de threshold tuning confirma que no existe un umbral único que haga '
    'cumplir los tres criterios simultáneamente para el modelo de Estrategia A, '
    'debido al compromiso (trade-off) inherente entre sensibilidad, especificidad y '
    'precisión. Al elevar el umbral para aumentar precisión y especificidad, la '
    'sensibilidad cae por debajo de 85 %.'
)
para(
    'En el contexto médico de NOR VISIÓN, se prioriza mantener la sensibilidad en '
    '87.59 % con el umbral 0.50: no detectar somnolencia en un médico fatigado '
    '(falso negativo) representa un riesgo mayor para la seguridad del paciente '
    'que emitir una alerta innecesaria (falso positivo). El ajuste fino del umbral '
    'operacional queda diferido a la validación con usuarios reales en el OE-06, '
    'donde se dispondrá de datos propios de NOR VISIÓN.'
)
para(
    'Las curvas ROC (AUC-ROC = 0.7942), matrices de confusión y curvas de entrenamiento '
    'generadas se adjuntan en el Anexo N.° 02.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# II. REFERENCIAS BIBLIOGRÁFICAS
# ═══════════════════════════════════════════════════════════════════════════════
h1('REFERENCIAS BIBLIOGRÁFICAS')

para(
    'Las referencias se presentan en estilo IEEE. Se recomienda importarlas a Zotero '
    'mediante DOI (Ctrl+Shift+I en Zotero → "Añadir elemento por identificador"). '
    'Las marcadas con [DOI pendiente] requieren búsqueda manual en Google Scholar '
    'o IEEE Xplore antes de incluirlas.'
)

refs = [
    ('[1]\t'
     'A. Azevedo y M. F. Santos, «KDD, SEMMA AND CRISP-DM: A parallel overview», '
     'en Proceedings of the IADIS European Conference on Data Mining, 2008, pp. 182-185. '
     '[DOI pendiente — buscar en Google Scholar: "KDD SEMMA CRISP-DM Azevedo Santos 2008"]'),

    ('[2]\t'
     'I. Nasri, «Driver Drowsiness Dataset (DDD)», Kaggle, 2022. [En línea]. '
     'Disponible en: https://www.kaggle.com/datasets/ismailnasri20/driver-drowsiness-dataset-ddd. '
     '[Accedido: 28 de abril de 2026]. [Citar como dataset en Zotero → tipo: Dataset]'),

    ('[3]\t'
     'A. Chowdhury et al., «Subject-independent drowsiness recognition from single-channel '
     'EEG with an interpretable CNN-LSTM model», Neurocomputing, vol. 552, p. 126581, 2023, '
     'doi: 10.1016/j.neucom.2023.126581. '
     '[Importar a Zotero con DOI: 10.1016/j.neucom.2023.126581]'),

    ('[4]\t'
     'S. Díaz-Santos, Ó. Cigala-Álvarez, E. Gonzalez-Sosa, P. Caballero-Gil, y '
     'C. Caballero-Gil, «Driver Identification and Detection of Drowsiness while Driving», '
     'Appl. Sci., vol. 14, n.° 6, p. 2603, mar. 2024, doi: 10.3390/app14062603. '
     '[Importar a Zotero con DOI: 10.3390/app14062603]'),

    ('[5]\t'
     'A. K. Pathak, A. K. Singh, P. Kumar, V. Bhatia, y O. Krejcar, «Real-time anti-sleep '
     'alert algorithm to prevent road accidents to ensure road safety», Front. Future Transp., '
     'vol. 6, p. 1545411, mar. 2025, doi: 10.3389/ffutr.2025.1545411. '
     '[Importar a Zotero con DOI: 10.3389/ffutr.2025.1545411]'),

    ('[6]\t'
     'F. H. Kamaru Zaman, K. M. Ng, y S. A. Che Abdullah, «Comparative Analysis of '
     'Vision Transformers and CNN Models for Driver Fatigue Classification», IIUM Eng. J., '
     'vol. 26, n.° 2, pp. 169-186, may 2025, doi: 10.31436/iiumej.v26i2.3488. '
     '[Importar a Zotero con DOI: 10.31436/iiumej.v26i2.3488]'),

    ('[7]\t'
     'S. A. El-Nabi et al., «Driver Drowsiness Detection Using Swin Transformer and '
     'Diffusion Models for Robust Image Denoising», IEEE Access, vol. 13, '
     'pp. 71880-71907, 2025, doi: 10.1109/ACCESS.2025.3561717. '
     '[Importar a Zotero con DOI: 10.1109/ACCESS.2025.3561717]'),

    ('[8]\t'
     'J. Wijsman, B. Grundlehner, J. Penders, y H. Hermens, «Trapezius muscle EMG as '
     'predictor of mental stress», ACM Trans. Embed. Comput. Syst., vol. 12, n.° 4, '
     'pp. 1-20, jun. 2013, doi: 10.1145/2485984.2485987. '
     '[Importar a Zotero con DOI: 10.1145/2485984.2485987]'),

    ('[9]\t'
     'J.-Y. Zhang, S.-L. Liu, Q.-M. Feng, J.-Q. Gao, y Q. Zhang, «Correlative Evaluation '
     'of Mental and Physical Workload of Laparoscopic Surgeons Based on Surface '
     'Electromyography and Eye-tracking Signals», Sci. Rep., vol. 7, n.° 1, p. 11095, '
     'sep. 2017, doi: 10.1038/s41598-017-11584-4. '
     '[Importar a Zotero con DOI: 10.1038/s41598-017-11584-4]'),
]

for ref in refs:
    para(ref)

# ═══════════════════════════════════════════════════════════════════════════════
# ANEXOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('ANEXOS')

h2('Anexo N.° 01. Carta de aceptación de NOR VISIÓN para la ejecución de la tesis')
para('[Adjuntar carta de aceptación firmada por la dirección de NOR VISIÓN.]')

h2('Anexo N.° 02. Gráficas de resultados del experimento')
para(
    'Las siguientes figuras fueron generadas durante la ejecución del notebook '
    'Modulo1_CNN_LSTM_ViT_v6_RUN_ALL_RESULTADOS.ipynb y se encuentran en la carpeta '
    'resultados/ del repositorio del proyecto:'
)
bullet('distribuciones_variables.png — Distribución de EAR y MAR por clase (alerta vs. somnoliento).')
bullet('matrices_confusion_dual.png — Matrices de confusión de los 6 modelos (3 arquitecturas × 2 estrategias).')
bullet('roc_curves_dual.png — Curvas ROC de los 6 modelos con AUC-ROC.')
bullet('training_curves.png — Curvas de pérdida y F1 durante el entrenamiento (Estrategia A).')
bullet('threshold_tuning.png — Análisis del umbral de decisión sobre el modelo ganador.')

# ─── Guardar ─────────────────────────────────────────────────────────────────
output_path = 'Pre_Informe_Tesis_TORRES_CABREJOS_2026.docx'
doc.save(output_path)
print(f'Documento guardado: {output_path}')
